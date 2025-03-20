import os
import random
import platform, shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph

FONT_FAMILY = [
    ("NinaCTT", 14.0),
    ("BetinaScriptC", 14.0),
    ("BirchCTT", 14.0),
    ("FreestyleScript", 18.0),
    ("MisterK", 22.0),
]


def replacer(txt: str):
    return txt.replace(
        "і", "i"
    ).replace(
        "І", "I"
    ).replace(
        "ї", "i"
    ).replace(
        "Ї", "i"
    )


def p_indent_l(ind: float, paragraph: Paragraph):
    """
    Відступи ліворуч
    :param ind:
    :param paragraph:
    :return:
    """
    format_par = paragraph.paragraph_format
    format_par.left_indent = Inches(ind)


def p_indent_r(ind: float, paragraph: Paragraph):
    """
    Відступи праворуч
    :param ind: A float multiplier to determine the scaling of the paragraph indent.
    :param paragraph: Paragraph object whose indentation is to be adjusted.
    :return: The updated paragraph indent after applying the scaling multiplier.
    :rtype: float
    """
    format_par = paragraph.paragraph_format
    format_par.right_indent = Inches(ind)


def p_indent_first_line(idn: float, paragraph: Paragraph):
    """
    Відступ першого рядка
    :param idn:
    :param paragraph:
    :return:
    """
    fromat_par = paragraph.paragraph_format
    fromat_par.first_line_indent = Inches(idn)


class GetDocument:
    def __init__(self, path_to_file: str):
        self.path_to_file = path_to_file
        self.doc = Document(self.path_to_file)
        self.font_type = random.choice(FONT_FAMILY)

    def change_font(self, paragraph: Paragraph):
        for pars in paragraph.runs:
            pars.text = replacer(pars.text)
            font_4 = pars.font
            font_4.size = Pt(self.font_type[1])
            font_4.name = self.font_type[0]


    def start_change_font(self):
        for paragraph in self.doc.paragraphs:
            self.change_font(paragraph)

        if not os.path.exists(f"report_new"):
            os.mkdir(f"report_new")
        self.doc.save(f"report_new/{self.path_to_file}.docx")


class CreateDocument:
    def __init__(
            self, first_name: str, last_name: str,
            military_rank: str,
            major_data: dict[str, str],
    ):
        self.first_name = first_name
        self.last_name = last_name
        self.military_rank = military_rank
        self.major_data = major_data
        self.doc = Document()
        self.font_type = random.choice(FONT_FAMILY)

    def set_paragraph_head(self, txt: str, left_indent=True):
        pargrph = self.doc.add_paragraph(txt)
        if left_indent:
            p_indent_l(2.7, pargrph)
        else:
            pargrph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_4 = pargrph.runs[0]
        font_4 = run_4.font
        font_4.size = Pt(self.font_type[1])
        font_4.name = self.font_type[0]

    def set_paragraph_main(self):
        text = self.doc.add_paragraph(replacer(self.major_data["text"]))
        run_4 = text.runs[0]
        font_4 = run_4.font
        font_4.size = Pt(self.font_type[1])
        font_4.name = self.font_type[0]

    def set_paragraph_date(self, txt: str):
        pargrph = self.doc.add_paragraph(txt)
        p_indent_l(4.5, pargrph)
        run_4 = pargrph.runs[0]
        font_4 = run_4.font
        font_4.size = Pt(self.font_type[1])
        font_4.name = self.font_type[0]

    def set_paragraph_from(self, txt: str):
        pargrph = self.doc.add_paragraph(txt)
        p_indent_r(4.0, pargrph)
        run_4 = pargrph.runs[0]
        font_4 = run_4.font
        font_4.size = Pt(self.font_type[1])
        font_4.name = self.font_type[0]

    #     Таблиця
    #     tb = self.doc.add_table(rows=1, cols=2)
    #     cell_1 = tb.cell(0, 0)
    #     cell_1.text = txt
    #     cell_2 = tb.cell(0, 1)
    #     cell_2.text = replacer(self.major_data["date"])
    #     cell_1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    #     cell_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #     cell_1.paragraphs[0].runs[0].font.size = Pt(self.font_type[1])
    #     cell_1.paragraphs[0].runs[0].font.name = self.font_type[0]
    #     cell_2.paragraphs[0].runs[0].font.size = Pt(self.font_type[1])
    #     cell_2.paragraphs[0].runs[0].font.name = self.font_type[0]
    #     cel

    def start_write_raports(self):
        self.set_paragraph_head(replacer("Командиру військової частини 3027"))
        self.set_paragraph_head(replacer(self.major_data["rank"] + "у"))
        self.set_paragraph_head(replacer(self.major_data["full_name"]))
        self.set_paragraph_head(replacer(f"{self.first_name} {self.last_name}"))
        self.doc.add_paragraph("\n")
        self.set_paragraph_head("Рапорт", left_indent=False)
        self.set_paragraph_main()

        self.doc.add_paragraph("\n")
        self.set_paragraph_from(
            replacer(
                f"{self.military_rank} {self.first_name} {self.last_name}"
            )
        )
        self.set_paragraph_date(replacer(self.major_data["date"]))
        self.doc.save(f"{self.first_name}_{self.last_name}_рапорт.docx")


class PreviousPrepared:
    """
    Поідготовка даних для зчитування та запису рапортів
    Файл з іменами вс
    Шрифти та куди їх копіювати в залежності від ОС
    """

    def __init__(self):
        self.file_names = r"names.txt"
        self.pathes_systems = {
            "Windows": r"C:\Windows\Fonts",
            "Linux": r"/usr/share/fonts",
            "Darwin": r"/Library/Fonts",
        }

    def copy_fonts(self) -> None:
        """
        Копіює шрифти із директорії fonts в належну директорію ОС
        :return:
        """

        os_name = platform.system()
        files_list = os.listdir(r"fonts")
        for x in range(len(files_list)):
            if not os.path.exists(
                    self.pathes_systems[os_name] + "/" + files_list[x]
            ) and not os.path.isfile(
                self.pathes_systems[os_name] + "/" + files_list[x]
            ):
                shutil.copy(r"fonts/{}".format(files_list[x]), self.pathes_systems[os_name])
            else:
                continue

    def open_get_names(self) -> list[str]:
        """
        Відкриває текстовий файл з іменами в/с та формує з них список
        для подальшого використання в формуванні *.docx файлів
        :return:
        """
        with open(self.file_names, "r", encoding="utf-8") as fl:
            content = fl.read()
            array_content = content.split("\n")
            return array_content


class GetterFiles:
    def __init__(self, path_files: str, cd: str):
        self.path_files = path_files
        self.cd = cd
        if self.cd:
            os.chdir(self.cd)
        else:
            os.chdir(self.path_files)

    def get_docx_files(self) -> list[str] | bool:
        """
        Отримує списки ворд файлів в директорії
        та повертає їх список
        інакше повертає None
        :return:
        """
        docx_files = []
        if self.path_files:
            files_list = os.listdir(self.path_files)
            if len(files_list) > 0:
                for x in range(len(files_list)):
                    if files_list[x].endswith(".docx") or files_list[x].endswith(".DOCX"):
                        docx_files.append(files_list[x])
                    else:
                        continue
            return docx_files
        else:
            return False


def run(data: dict, switcher: bool = True, path: str = None):
    """
    Запускає цикл перебору імен
    та запускає метод класу CreateDocument для формування *.docx файлів
    або викликає GetterFiles і зміює шрифт у копіях файлів
    :return:
    """
    helper = PreviousPrepared()
    helper.copy_fonts()
    if not switcher:
        array_militaries = helper.open_get_names()
        for x in range(len(array_militaries)):
            first = array_militaries[x].split()[0].strip()
            last = array_militaries[x].split()[1].strip()
            rank = array_militaries[x].split()[2].strip()
            my_doc = CreateDocument(
                first_name=first,
                last_name=last,
                military_rank=rank,
                major_data=data,
            )
            my_doc.start_write_raports()
    else:
        print("Strart copy")
        gf = GetterFiles(data["path_to_directory"], cd=path)
        files = gf.get_docx_files()
        print(f"Знайдено {len(files)} файлів *.docx")
        if files:
            for y in range(len(files)):
                print(f"Копіюємо {files[y]}")
                fl_doc = GetDocument(files[y])
                fl_doc.start_change_font()
                print(f"Завершено копіювання {files[y]}")
        print("End copy")
